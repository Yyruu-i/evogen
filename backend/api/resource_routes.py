"""资源库 REST API — 技能 CRUD / 导入导出 / 工具注册 / 经验导出.

端点前缀：/api/v1/resource
"""

import json
import logging
import os
import re
import shutil
import uuid
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Optional
from zipfile import ZipFile

import yaml
from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from backend.api.tools_routes import _static_tool_list as _hermes_static_tools
from backend.auth.dependencies import get_current_user as get_user_id

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/resource", tags=["resource"])

# ── 配置 ──
# 技能目录：扫描所有 Hermes profile 的 skills 子目录 + 系统级 ~/.hermes/skills
def _resolve_skills_dirs() -> list[Path]:
    """解析所有可能的技能目录，去重排序."""
    candidates = []

    # 1. 所有 Hermes profile 的 skills 目录（glob /root/.hermes/profiles/*/skills）
    profiles_base = Path("/root/.hermes/profiles")
    if profiles_base.exists():
        for profile_dir in sorted(profiles_base.iterdir()):
            if profile_dir.is_dir():
                sd = profile_dir / "skills"
                if sd.is_dir():
                    candidates.append(sd)

    # 2. 当前 HERMES_HOME （如果有且未在 profiles 中覆盖）
    hh = os.environ.get("HERMES_HOME")
    if hh:
        p = Path(hh) / "skills"
        if p not in candidates:
            candidates.append(p)

    # 3. 系统级 ~/.hermes/skills
    candidates.append(Path("/root/.hermes/skills"))
    candidates.append(Path(os.path.expanduser("~/.hermes/skills")))

    # 去重
    seen = set()
    result = []
    for p in candidates:
        if p not in seen:
            seen.add(p)
            result.append(p)
    return result

_SKILLS_DIRS: list[Path] = _resolve_skills_dirs()
# 主写入目录：必须使用当前 HERMES_HOME 的 skills（即 backend profile），
# 不能取 _SKILLS_DIRS[0]（按字母序可能是 architect 等只读 profile）
def _get_write_dir() -> Path:
    """每次调用动态计算写入目录 — 避免模块导入时 HERMES_HOME 未设置."""
    hh = os.environ.get("HERMES_HOME")
    if hh:
        d = Path(hh) / "skills"
    else:
        # 回退：在 _SKILLS_DIRS 中找路径含 "backend" 的
        for sd in _SKILLS_DIRS:
            if "backend" in str(sd).lower() and sd.exists():
                d = sd
                break
        else:
            # 再回退：显式构造 backend profile 路径
            d = Path(os.path.expanduser("~/.hermes/profiles/backend/skills"))
            for sd in _SKILLS_DIRS:
                if sd.exists():
                    d = sd
                    break
    d.mkdir(parents=True, exist_ok=True)
    return d

_USER_DIR_RE = re.compile(
    r"^(default|[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}|user_[a-z0-9_]+)$"
)
_HERMES_HOME = os.environ.get("HERMES_HOME", os.path.expanduser("~/.hermes"))
_TOOL_REGISTRY_PATH = Path(_HERMES_HOME) / "tools_registry.json"

# 确保工具注册表目录存在
_TOOL_REGISTRY_PATH.parent.mkdir(parents=True, exist_ok=True)

# ─────────────────────────────────────────────────────
# 工具注册表（本地 JSON 文件）
# ─────────────────────────────────────────────────────


def _load_tools_registry() -> dict:
    """加载工具注册表."""
    if _TOOL_REGISTRY_PATH.exists():
        try:
            return json.loads(_TOOL_REGISTRY_PATH.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            logger.warning("Corrupt tools_registry.json, resetting")
    return {"tools": {}}


def _save_tools_registry(registry: dict) -> None:
    """保存工具注册表."""
    _TOOL_REGISTRY_PATH.write_text(
        json.dumps(registry, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def _write_file_sync(path: Path, content: str) -> None:
    """写入文件并强制 fsync 到磁盘，确保数据真实落盘."""
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
        f.flush()
        os.fsync(f.fileno())


# ════════════════════════════════════════════════════════


class CreateSkillRequest(BaseModel):
    name: str = Field(..., description="技能名称")
    description: str = Field("", description="技能描述")
    content: str = Field(..., description="SKILL.md 内容")
    category: str = Field("", description="分类目录名")


class UpdateSkillRequest(BaseModel):
    name: Optional[str] = Field(None, description="技能名称")
    description: Optional[str] = Field(None, description="技能描述")
    content: Optional[str] = Field(None, description="SKILL.md 内容")
    category: Optional[str] = Field(None, description="分类目录名")


class ExportSkillsRequest(BaseModel):
    skill_ids: list[str] = Field(..., description="要导出的技能 ID 列表")


class ImportSkillsRequest(BaseModel):
    """用于解析导入的 JSON 元数据（文件上传走 UploadFile）."""


class RegisterToolRequest(BaseModel):
    name: str = Field(..., description="工具名称")
    description: str = Field(..., description="工具描述")
    endpoint: str = Field("", description="工具端点/命令")
    category: str = Field("", description="工具分类")
    parameters: dict = Field(default_factory=dict, description="参数 schema")


class UpdateToolRequest(BaseModel):
    name: Optional[str] = Field(None, description="工具名称")
    description: Optional[str] = Field(None, description="工具描述")
    endpoint: Optional[str] = Field(None, description="工具端点/命令")
    category: Optional[str] = Field(None, description="工具分类")
    parameters: Optional[dict] = Field(None, description="参数 schema")


class ExportExperienceRequest(BaseModel):
    trajectory_ids: Optional[list[str]] = Field(None, description="要导出的轨迹 ID（None=全部）")
    format: str = Field("json", description="导出格式: json / zip")


# ════════════════════════════════════════════════════════
# 辅助函数
# ════════════════════════════════════════════════════════


def _list_all_skill_ids() -> list[str]:
    """列出所有已安装的技能 ID（目录名），扫描所有技能目录."""
    ids = set()
    for sd in _SKILLS_DIRS:
        if not sd.exists():
            continue
        for md in sd.rglob("SKILL.md"):
            if md.is_file():
                ids.add(md.parent.name)
    return sorted(ids)


def _skill_dir(skill_id: str) -> Path:
    """返回技能目录路径（扫描所有技能目录）.

    搜索策略：遍历所有 _SKILLS_DIRS，rglob 扫描 SKILL.md，匹配父目录名。
    """
    for sd in _SKILLS_DIRS:
        if not sd.exists():
            continue
        for md in sd.rglob("SKILL.md"):
            if md.is_file() and md.parent.name == skill_id:
                return md.parent
    raise HTTPException(status_code=404, detail={"ok": False, "error": f"Skill not found: {skill_id}"})


def _parse_frontmatter(md_content: str) -> dict:
    """从 SKILL.md 内容中解析 YAML frontmatter."""
    if not md_content.startswith("---"):
        return {}
    parts = md_content.split("---", 2)
    if len(parts) < 3:
        return {}
    try:
        return yaml.safe_load(parts[1]) or {}
    except yaml.YAMLError:
        return {}


def _build_skill_md(name: str, description: str, category: str, content: str, user_id: str = "") -> str:
    """构建带 frontmatter 的 SKILL.md."""
    fm = {
        "name": name,
        "description": description,
        "version": "1.0.0",
        "scope": "user",
    }
    if category:
        fm["category"] = category
    if user_id:
        fm["user_id"] = user_id
    fm_yaml = yaml.dump(fm, allow_unicode=True, default_flow_style=False).strip()
    return f"---\n{fm_yaml}\n---\n\n{content}"


def _skill_to_dict(skill_id: str, md_path: Path, scope: str = "builtin") -> dict:
    """将技能文件转为 API 友好字典 — scope 优先从前置元数据读取."""
    content = md_path.read_text(encoding="utf-8")
    fm = _parse_frontmatter(content)
    mtime = md_path.stat().st_mtime
    created_at = datetime.fromtimestamp(mtime, tz=timezone.utc).isoformat()
    # 优先从前置元数据读取 scope，没有则回退到参数
    resolved_scope = fm.get("scope", scope)
    return {
        "id": skill_id,
        "name": fm.get("name", skill_id),
        "description": fm.get("description", ""),
        "category": fm.get("category", ""),
        "version": fm.get("version", "1.0.0"),
        "scope": resolved_scope,
        "user_id": fm.get("user_id", ""),
        "created_at": created_at,
        "file_size": len(content),
    }


# ════════════════════════════════════════════════════════
# GET /api/v1/resource/skills — 列出技能
# ════════════════════════════════════════════════════════


@router.get("/skills")
async def list_resource_skills(user_id: str = Depends(get_user_id)):
    """列出所有技能目录下的技能（多源扫描，去重，按用户隔离）."""
    seen: set[str] = set()
    skills = []
    for sd in _SKILLS_DIRS:
        if not sd.exists():
            continue
        for md in sd.rglob("SKILL.md"):
            if not md.is_file():
                continue
            skill_id = md.parent.name
            if skill_id in seen:
                continue

            # 读取前置元数据确定 scope 和 user_id
            try:
                fm = _parse_frontmatter(md.read_text(encoding="utf-8"))
            except Exception:
                fm = {}

            fm_scope = fm.get("scope", "")
            fm_user_id = fm.get("user_id", "")

            if fm_scope:
                # 有明确 scope 元数据 — 直接使用，按 user_id 过滤
                if fm_scope == "user" and fm_user_id and fm_user_id != user_id:
                    continue
                scope = fm_scope
            else:
                # 回退到路径推断（兼容旧数据）
                # 仅当路径第一级看起来像用户目录时才做隔离判断
                is_user_skill = False
                try:
                    rel = md.parent.relative_to(sd)
                    parts = rel.parts
                    if len(parts) >= 2:
                        if parts[0] == user_id:
                            is_user_skill = True
                        elif _USER_DIR_RE.match(parts[0]):
                            continue  # 其他用户的目录
                except ValueError:
                    pass
                scope = "user" if is_user_skill else "builtin"

            seen.add(skill_id)
            skills.append(_skill_to_dict(skill_id, md, scope=scope))
    return {"ok": True, "data": {"skills": skills, "total": len(skills)}}


# ════════════════════════════════════════════════════════
# GET /api/v1/resource/skills/{skill_id} — 获取单个技能
# ════════════════════════════════════════════════════════


@router.get("/skills/{skill_id}")
async def get_resource_skill(skill_id: str, user_id: str = Depends(get_user_id)):
    """获取单个技能的完整内容."""
    skill_dir = _skill_dir(skill_id)
    md_file = skill_dir / "SKILL.md"
    content = md_file.read_text(encoding="utf-8")
    # 从前置元数据读取 scope，回退到路径推断
    fm = _parse_frontmatter(content)
    scope = fm.get("scope", "")
    if not scope:
        # 回退到路径推断（兼容旧数据）
        scope = "builtin"
        try:
            write_dir = _get_write_dir()
            md_file.resolve().relative_to((write_dir / user_id).resolve())
            scope = "user"
        except ValueError:
            pass
    return {"ok": True, "data": {"skill": _skill_to_dict(skill_id, md_file, scope=scope), "content": content}}


# ════════════════════════════════════════════════════════
# POST /api/v1/resource/skills  — 创建技能
# ════════════════════════════════════════════════════════


@router.post("/skills", status_code=201)
async def create_resource_skill(request: CreateSkillRequest, user_id: str = Depends(get_user_id)):
    """创建新技能 — 在用户专属目录下创建 ~/.hermes/skills/{user_id}/SKILL.md."""
    # 生成安全 ID
    skill_id = request.name.lower().replace(" ", "-").replace("_", "-")

    # 确定目标目录（用户隔离：所有自定义技能写入系统级 skills 目录的 {user_id} 子目录）
    # _SKILLS_DIRS[-1] 始终是 ~/.hermes/skills（排在所有 profile skills 之后）
    user_skills_base = _SKILLS_DIRS[-1] / user_id
    if request.category:
        target_dir = user_skills_base / request.category / skill_id
    else:
        target_dir = user_skills_base / skill_id

    if target_dir.exists():
        raise HTTPException(
            status_code=409,
            detail={"ok": False, "error": f"Skill already exists: {skill_id}"},
        )

    target_dir.mkdir(parents=True, exist_ok=False)

    md_content = _build_skill_md(
        name=request.name,
        description=request.description,
        category=request.category,
        content=request.content,
        user_id=user_id,
    )

    md_file = target_dir / "SKILL.md"
    _write_file_sync(md_file, md_content)

    logger.info(f"Skill created: {skill_id} at {target_dir}")
    return {
        "ok": True,
        "data": _skill_to_dict(skill_id, md_file, scope="user"),
    }


# ════════════════════════════════════════════════════════
# PUT /api/v1/resource/skills/{skill_id} — 更新技能
# ════════════════════════════════════════════════════════


@router.put("/skills/{skill_id}")
async def update_resource_skill(skill_id: str, request: UpdateSkillRequest, user_id: str = Depends(get_user_id)):
    """更新技能 — 仅允许编辑用户自定义技能（scope=user），内置技能不可编辑.

    如果更新 category，会移动目录。
    """
    skill_dir = _skill_dir(skill_id)
    md_file = skill_dir / "SKILL.md"

    # 读取现有内容 + frontmatter
    current_content = md_file.read_text(encoding="utf-8")
    fm = _parse_frontmatter(current_content)

    # ── 安全检查：内置技能不可编辑 ──
    scope = fm.get("scope", "")
    if scope == "builtin":
        raise HTTPException(
            status_code=403,
            detail={"ok": False, "error": "内置技能（scope=builtin）不可编辑，仅用户自定义技能可编辑"},
        )
    if not scope:
        # 回退到路径推断（兼容旧数据）
        user_skills_base = _get_write_dir() / user_id
        try:
            skill_dir.resolve().relative_to(user_skills_base.resolve())
        except ValueError:
            raise HTTPException(
                status_code=403,
                detail={"ok": False, "error": "内置技能不可编辑，仅用户自定义技能可编辑"},
            )
    # 检查 user_id 归属
    fm_user_id = fm.get("user_id", "")
    if fm_user_id and fm_user_id != user_id:
        raise HTTPException(
            status_code=403,
            detail={"ok": False, "error": "无权编辑其他用户的技能"},
        )

    # 提取 body（frontmatter 之后的内容，保留原始空白）
    body = ""
    if current_content.startswith("---"):
        parts = current_content.split("---", 2)
        if len(parts) >= 3:
            body = parts[2]  # 不去 strip，保留正文原有格式

    new_name = request.name if request.name is not None else fm.get("name", skill_id)
    new_desc = request.description if request.description is not None else fm.get("description", "")
    new_category = request.category if request.category is not None else fm.get("category", "")
    new_body = request.content if request.content is not None else body

    # 处理改名 / 改分类 → 目录移动
    new_skill_id = new_name.lower().replace(" ", "-").replace("_", "-")
    write_dir = _get_write_dir()
    user_skills_base = write_dir / user_id
    if new_category:
        new_target_dir = user_skills_base / new_category / new_skill_id
    else:
        new_target_dir = user_skills_base / new_skill_id

    if new_target_dir != skill_dir:
        if new_target_dir.exists():
            raise HTTPException(
                status_code=409,
                detail={"ok": False, "error": f"Target skill already exists: {new_skill_id}"},
            )
        new_target_dir.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(skill_dir), str(new_target_dir))
        md_file = new_target_dir / "SKILL.md"
        skill_id = new_skill_id

    # 写入新内容（使用 fsync 确保落盘）
    new_md = _build_skill_md(new_name, new_desc, new_category, new_body, user_id=user_id)
    _write_file_sync(md_file, new_md)

    logger.info(f"Skill updated: {skill_id}")
    return {"ok": True, "data": _skill_to_dict(skill_id, md_file, scope="user")}


# ════════════════════════════════════════════════════════
# DELETE /api/v1/resource/skills/{skill_id} — 删除技能
# ════════════════════════════════════════════════════════


class BatchDeleteRequest(BaseModel):
    skill_ids: list[str] = Field(..., description="要删除的技能 ID 列表")


def _delete_skill_by_id(skill_id: str, user_id: str = "default") -> bool:
    """删除技能文件（调用方已通过 scope 检查）。返回是否成功."""
    try:
        skill_dir = _skill_dir(skill_id)
    except HTTPException:
        return False
    shutil.rmtree(str(skill_dir))
    logger.info(f"Skill deleted: {skill_id} from user={user_id}")
    return True


@router.delete("/skills/{skill_id}")
async def delete_resource_skill(skill_id: str, user_id: str = Depends(get_user_id)):
    """删除技能 — 移除整个技能目录（仅限用户自有技能）."""
    # 先做 scope 检查
    skill_dir = _skill_dir(skill_id)
    md_file = skill_dir / "SKILL.md"
    try:
        fm = _parse_frontmatter(md_file.read_text(encoding="utf-8"))
    except Exception:
        fm = {}
    scope = fm.get("scope", "")
    if scope == "builtin":
        raise HTTPException(status_code=403, detail={"ok": False, "error": "内置技能（scope=builtin）不可删除"})
    if not scope:
        # 回退到路径推断（兼容旧数据）
        user_skills_base = _get_write_dir() / user_id
        try:
            skill_dir.resolve().relative_to(user_skills_base.resolve())
        except ValueError:
            raise HTTPException(status_code=403, detail={"ok": False, "error": "内置技能不可删除"})
    fm_user_id = fm.get("user_id", "")
    if fm_user_id and fm_user_id != user_id:
        raise HTTPException(status_code=403, detail={"ok": False, "error": "无权删除其他用户的技能"})

    if not _delete_skill_by_id(skill_id, user_id=user_id):
        raise HTTPException(status_code=404, detail={"ok": False, "error": f"Skill not found: {skill_id}"})
    return {"ok": True, "data": {"deleted": skill_id}}


@router.post("/skills/batch-delete")
async def batch_delete_resource_skills(request: BatchDeleteRequest, user_id: str = Depends(get_user_id)):
    """批量删除技能.

    请求体: {"skill_ids": ["skill-a", "skill-b", ...]}
    返回: {"ok": true, "data": {"deleted": [...], "not_found": [...], "forbidden": [...], "total": N}}
    """
    deleted: list[str] = []
    not_found: list[str] = []
    forbidden: list[str] = []
    for sid in request.skill_ids:
        # 先做 scope 检查
        try:
            skill_dir = _skill_dir(sid)
            md_file = skill_dir / "SKILL.md"
            try:
                fm = _parse_frontmatter(md_file.read_text(encoding="utf-8"))
            except Exception:
                fm = {}
            scope = fm.get("scope", "")
            if scope == "builtin":
                forbidden.append(sid)
                continue
            if not scope:
                user_skills_base = _get_write_dir() / user_id
                try:
                    skill_dir.resolve().relative_to(user_skills_base.resolve())
                except ValueError:
                    forbidden.append(sid)
                    continue
            fm_user_id = fm.get("user_id", "")
            if fm_user_id and fm_user_id != user_id:
                not_found.append(sid)
                continue
            if _delete_skill_by_id(sid, user_id=user_id):
                deleted.append(sid)
            else:
                not_found.append(sid)
        except HTTPException:
            not_found.append(sid)
    logger.info(f"Batch delete: {len(deleted)} deleted, {len(not_found)} not found, {len(forbidden)} forbidden")
    result = {"deleted": deleted, "not_found": not_found, "total": len(request.skill_ids)}
    if forbidden:
        result["forbidden"] = forbidden
    return {"ok": True, "data": result}


# ════════════════════════════════════════════════════════
# POST /api/v1/resource/skills/export — 导出技能
# ════════════════════════════════════════════════════════


@router.post("/skills/export")
async def export_resource_skills(request: ExportSkillsRequest):
    """导出技能 — 单个/批量打包为 ZIP 下载.

    请求体: {"skill_ids": ["skill-a", "skill-b"]}
    返回: ZIP 文件流
    """
    if not request.skill_ids:
        raise HTTPException(
            status_code=400,
            detail={"ok": False, "error": "skill_ids is required"},
        )

    # 验证所有技能存在
    skill_dirs: list[tuple[str, Path]] = []
    for sid in request.skill_ids:
        skill_dirs.append((sid, _skill_dir(sid)))

    # 打包 ZIP
    zip_buffer = BytesIO()
    with ZipFile(zip_buffer, "w") as zf:
        for sid, sdir in skill_dirs:
            for file_path in sdir.rglob("*"):
                if file_path.is_file():
                    arcname = f"{sid}/{file_path.relative_to(sdir)}"
                    zf.write(file_path, arcname)

    zip_buffer.seek(0)
    filename = f"skills-export-{len(request.skill_ids)}-{uuid.uuid4().hex[:8]}.zip"

    logger.info(f"Skills exported: {request.skill_ids}")
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ════════════════════════════════════════════════════════
# POST /api/v1/resource/skills/import — 导入技能
# ════════════════════════════════════════════════════════


@router.post("/skills/import")
async def import_resource_skills(file: UploadFile, user_id: str = Depends(get_user_id)):
    """导入技能 — 上传 .md 或 .json(批量) 或 .zip，写入用户专属目录."""
    if not file.filename:
        raise HTTPException(status_code=400, detail={"ok": False, "error": "No file provided"})

    raw = await file.read()
    filename = file.filename.lower()
    imported: list[str] = []
    user_base = _get_write_dir() / user_id

    if filename.endswith(".zip"):
        zip_buffer = BytesIO(raw)
        with ZipFile(zip_buffer, "r") as zf:
            for member in zf.namelist():
                if member.endswith("/") or "__MACOSX" in member:
                    continue
                parts = member.split("/")
                if len(parts) < 2:
                    continue
                skill_id = parts[0]
                rel_path = "/".join(parts[1:])
                target = user_base / skill_id / rel_path
                target.parent.mkdir(parents=True, exist_ok=True)
                target.write_bytes(zf.read(member))
                if skill_id not in imported:
                    imported.append(skill_id)

    elif filename.endswith(".json"):
        data = json.loads(raw.decode("utf-8"))
        if isinstance(data, dict):
            data = [data]
        for item in data:
            skill_id = item["name"].lower().replace(" ", "-").replace("_", "-")
            category = item.get("category", "")
            if category:
                target_dir = user_base / category / skill_id
            else:
                target_dir = user_base / skill_id
            target_dir.mkdir(parents=True, exist_ok=True)
            md_content = _build_skill_md(
                name=item["name"],
                description=item.get("description", ""),
                category=category,
                content=item.get("content", ""),
                user_id=user_id,
            )
            _write_file_sync(target_dir / "SKILL.md", md_content)
            imported.append(skill_id)

    elif filename.endswith(".md"):
        content = raw.decode("utf-8")
        fm = _parse_frontmatter(content)
        name = fm.get("name", file.filename.replace(".md", ""))
        skill_id = name.lower().replace(" ", "-").replace("_", "-")
        category = fm.get("category", "")
        if category:
            target_dir = user_base / category / skill_id
        else:
            target_dir = user_base / skill_id
        target_dir.mkdir(parents=True, exist_ok=True)
        _write_file_sync(target_dir / "SKILL.md", content)
        imported.append(skill_id)

    else:
        raise HTTPException(
            status_code=400,
            detail={"ok": False, "error": f"Unsupported file type: {file.filename}. Supported: .md, .json, .zip"},
        )

    logger.info(f"Skills imported: {imported}")
    return {"ok": True, "data": {"imported": imported, "count": len(imported)}}


# ════════════════════════════════════════════════════════
# POST /api/v1/resource/skills/import/json — JSON 批量导入
# ════════════════════════════════════════════════════════


@router.post("/skills/import/json")
async def import_resource_skills_json(request: dict, user_id: str = Depends(get_user_id)):
    """通过 JSON body 批量导入技能到用户专属目录."""
    items = request.get("skills", [])
    if not items:
        raise HTTPException(
            status_code=400,
            detail={"ok": False, "error": "Missing 'skills' array in request body"},
        )

    imported: list[str] = []
    user_base = _get_write_dir() / user_id
    for item in items:
        name = item.get("name", "unnamed")
        skill_id = name.lower().replace(" ", "-").replace("_", "-")
        category = item.get("category", "")
        if category:
            target_dir = user_base / category / skill_id
        else:
            target_dir = user_base / skill_id
        target_dir.mkdir(parents=True, exist_ok=True)
        md_content = _build_skill_md(
            name=name,
            description=item.get("description", ""),
            category=category,
            content=item.get("content", ""),
            user_id=user_id,
        )
        _write_file_sync(target_dir / "SKILL.md", md_content)
        imported.append(skill_id)

    return {"ok": True, "data": {"imported": imported, "count": len(imported)}}


# ════════════════════════════════════════════════════════
# GET /api/v1/resource/tools — 列出工具
# ════════════════════════════════════════════════════════

# 从 tools_routes 获取内置工具列表并统一包装格式
_TOOLSET_CATEGORY: dict[str, str] = {
    "browser": "浏览器",
    "terminal": "终端",
    "file": "文件",
    "web": "联网搜索",
    "memory": "记忆",
    "session_search": "搜索",
    "delegation": "代理",
    "cronjob": "计划任务",
    "code_execution": "代码执行",
    "clarify": "AI",
    "messaging": "消息",
    "todo": "任务",
    "vision": "AI",
    "tts": "AI",
    "skills": "技能",
    "feishu_doc": "飞书",
    "feishu_drive": "飞书",
}


def _build_builtin_tools() -> list[dict]:
    """从 tools_routes._static_tool_list() 生成统一格式的内置工具列表."""
    raw = _hermes_static_tools()
    result: list[dict] = []
    for t in raw:
        name = t["name"]
        toolset = t.get("toolset", "")
        category = _TOOLSET_CATEGORY.get(toolset, toolset)
        # 将 parameters list 转成 {name: type} 字典格式
        params: dict[str, str] = {}
        for p in t.get("parameters", []):
            pname = p.get("name", "")
            ptype = p.get("type", "string")
            if pname:
                params[pname] = ptype
        result.append({
            "id": f"builtin_{name}",
            "name": name,
            "description": t.get("description", ""),
            "endpoint": f"/api/v1/{toolset}/{name.split('_')[-1] if '_' in name else name}",
            "category": category,
            "parameters": params,
            "scope": "builtin",
            "created_at": "2025-01-01T00:00:00+00:00",
        })
    return result


@router.get("/tools")
async def list_resource_tools(user_id: str = Depends(get_user_id)):
    """列出工具：内置工具（全局共享）+ 用户自定义工具（隔离）."""
    builtin_tools = _build_builtin_tools()
    registry = _load_tools_registry()
    all_tools = registry.get("tools", {})
    user_tools = list(all_tools.get(user_id, {}).values())
    # 标记用户工具 scope
    for t in user_tools:
        t["scope"] = "user"
    # 内置工具 + 用户工具合并
    return {"ok": True, "data": {"tools": builtin_tools + user_tools, "total": len(builtin_tools) + len(user_tools)}}


# ════════════════════════════════════════════════════════
# POST /api/v1/resource/tools — 注册新工具
# ════════════════════════════════════════════════════════


@router.post("/tools", status_code=201)
async def register_resource_tool(request: RegisterToolRequest, user_id: str = Depends(get_user_id)):
    """注册新工具到用户专属工具注册表."""
    registry = _load_tools_registry()
    all_tools = registry.setdefault("tools", {})
    user_tools = all_tools.setdefault(user_id, {})

    tool_id = str(uuid.uuid4())
    tool_entry = {
        "id": tool_id,
        "name": request.name,
        "description": request.description,
        "endpoint": request.endpoint,
        "category": request.category,
        "parameters": request.parameters,
        "user_id": user_id,
        "scope": "user",
        "created_at": datetime.now(timezone.utc).isoformat(),
    }

    user_tools[tool_id] = tool_entry
    _save_tools_registry(registry)

    logger.info(f"Tool registered: {request.name} (id={tool_id}, user={user_id})")
    return {"ok": True, "data": tool_entry}


# ════════════════════════════════════════════════════════
# PUT /api/v1/resource/tools/{tool_id} — 编辑工具
# ════════════════════════════════════════════════════════


@router.put("/tools/{tool_id}")
async def update_resource_tool(tool_id: str, request: UpdateToolRequest, user_id: str = Depends(get_user_id)):
    """编辑工具 — 仅允许编辑用户自定义工具（scope=user），内置工具不可编辑."""
    # 拒绝编辑内置工具（scope=builtin）
    if tool_id.startswith("builtin_"):
        raise HTTPException(
            status_code=403,
            detail={"ok": False, "error": "内置工具（scope=builtin）不可编辑，仅用户自定义工具可编辑"},
        )

    registry = _load_tools_registry()
    all_tools = registry.get("tools", {})
    user_tools = all_tools.get(user_id, {})

    if tool_id not in user_tools:
        raise HTTPException(
            status_code=404,
            detail={"ok": False, "error": f"Tool not found: {tool_id}"},
        )

    tool = user_tools[tool_id]
    if request.name is not None:
        tool["name"] = request.name
    if request.description is not None:
        tool["description"] = request.description
    if request.endpoint is not None:
        tool["endpoint"] = request.endpoint
    if request.category is not None:
        tool["category"] = request.category
    if request.parameters is not None:
        tool["parameters"] = request.parameters

    _save_tools_registry(registry)

    logger.info(f"Tool updated: {tool['name']} (id={tool_id}, user={user_id})")
    return {"ok": True, "data": tool}


# ════════════════════════════════════════════════════════
# DELETE /api/v1/resource/tools/{tool_id} — 删除工具
# ════════════════════════════════════════════════════════


@router.delete("/tools/{tool_id}")
async def delete_resource_tool(tool_id: str, user_id: str = Depends(get_user_id)):
    """删除工具 — 仅允许删除用户自定义工具，内置工具不可删除."""
    # 拒绝删除内置工具（scope=builtin）
    if tool_id.startswith("builtin_"):
        raise HTTPException(
            status_code=403,
            detail={"ok": False, "error": "内置工具（scope=builtin）不可删除"},
        )

    registry = _load_tools_registry()
    all_tools = registry.get("tools", {})
    user_tools = all_tools.get(user_id, {})

    if tool_id not in user_tools:
        raise HTTPException(
            status_code=404,
            detail={"ok": False, "error": f"Tool not found: {tool_id}"},
        )

    deleted = user_tools.pop(tool_id)
    _save_tools_registry(registry)

    logger.info(f"Tool deleted: {deleted['name']} (id={tool_id}, user={user_id})")
    return {"ok": True, "data": {"deleted": tool_id}}


# ════════════════════════════════════════════════════════
# POST /api/v1/resource/experience/export — 经验导出
# ════════════════════════════════════════════════════════


@router.post("/experience/export")
async def export_resource_experience(request: ExportExperienceRequest):
    """导出经验轨迹为 JSON 或 ZIP.

    请求体: {"trajectory_ids": ["id1","id2"], "format": "json"}
    如果 trajectory_ids 为空，导出全部。
    """
    from backend.experience.recorder import get_recorder

    recorder = get_recorder()

    # 确定要导出的轨迹
    if request.trajectory_ids:
        trajectory_ids = request.trajectory_ids
    else:
        summaries = recorder.list_trajectories(limit=10000)
        trajectory_ids = [s.id for s in summaries]

    if not trajectory_ids:
        raise HTTPException(
            status_code=404,
            detail={"ok": False, "error": "No trajectories found"},
        )

    # 获取详情
    details = []
    for tid in trajectory_ids:
        detail = recorder.get_trajectory(tid)
        if detail is None:
            continue
        details.append({
            "id": detail.id,
            "session_id": detail.session_id,
            "session_title": detail.session_title,
            "turns": [
                {
                    "turn_index": t.turn_index,
                    "llm_response_chunk": t.llm_response_chunk,
                    "token_usage": t.token_usage,
                }
                for t in detail.turns
            ],
            "outcome": {
                "success": detail.outcome.success,
                "total_tokens": detail.outcome.total_tokens,
                "wall_time_ms": detail.outcome.wall_time_ms,
                "user_cancelled": detail.outcome.user_cancelled,
            },
            "created_at": detail.created_at,
            "feedback": [
                {
                    "id": fb.id,
                    "rating": fb.rating,
                    "note": fb.note,
                    "status": fb.status,
                    "created_at": fb.created_at,
                }
                for fb in detail.feedback
            ],
        })

    if request.format == "json":
        export_json = json.dumps({"trajectories": details}, ensure_ascii=False, indent=2)
        return {
            "ok": True,
            "data": {
                "exported_count": len(details),
                "json": export_json,
            },
        }

    elif request.format == "zip":
        zip_buffer = BytesIO()
        with ZipFile(zip_buffer, "w") as zf:
            # 整体 JSON
            zf.writestr(
                "trajectories.json",
                json.dumps({"trajectories": details}, ensure_ascii=False, indent=2),
            )
            # 每个轨迹一个文件
            for d in details:
                zf.writestr(
                    f"trajectories/{d['id']}.json",
                    json.dumps(d, ensure_ascii=False, indent=2),
                )

        zip_buffer.seek(0)
        filename = f"experience-export-{len(details)}-{uuid.uuid4().hex[:8]}.zip"
        return StreamingResponse(
            zip_buffer,
            media_type="application/zip",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    else:
        raise HTTPException(
            status_code=400,
            detail={"ok": False, "error": f"Unsupported format: {request.format}. Use json or zip"},
        )


# ── Word 文档导出 ──

class ExportDocxRequest(BaseModel):
    content: str = ""
    title: str = "EvoGen 文档"
    artifact_id: Optional[str] = None
    session_id: Optional[str] = None


@router.post("/export-docx")
async def export_docx(request: ExportDocxRequest):
    """将 Markdown 内容或制品导出为 Word 文档 (.docx).

    支持两种模式：
    1. content + title：直接传入 Markdown 内容
    2. artifact_id：从制品库读取内容
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    content = request.content

    # 如果提供了 artifact_id，从制品库读取内容
    if request.artifact_id and not content:
        try:
            # artifacts 存储在 messages 表或 artifacts 表
            from backend.db.connection import get_db
            db = get_db()
            row = db.execute(
                "SELECT content FROM messages WHERE id=? AND role='assistant'",
                (request.artifact_id,),
            ).fetchone()
            if row:
                # 尝试从消息内容中提取制品（代码块等）
                msg_content = row["content"]
                # 查找代码块
                code_blocks = re.findall(r"```(\w*)\n(.*?)```", msg_content, re.DOTALL)
                if code_blocks:
                    content = "\n\n".join(
                        f"## {lang or '代码'}\n\n{code.strip()}"
                        for lang, code in code_blocks
                    )
                else:
                    content = msg_content
            else:
                raise HTTPException(status_code=404, detail="制品未找到")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"读取制品失败: {str(e)[:200]}",
            )

    if not content.strip():
        raise HTTPException(status_code=400, detail="内容不能为空")

    # 生成 Word 文档
    doc = Document()
    doc.add_heading(request.title or "EvoGen 文档", level=0)

    # 解析 Markdown 风格格式
    for line in content.split("\n"):
        line_stripped = line.strip()

        # 标题
        if line_stripped.startswith("### "):
            doc.add_heading(line_stripped[4:], level=3)
        elif line_stripped.startswith("## "):
            doc.add_heading(line_stripped[3:], level=2)
        elif line_stripped.startswith("# "):
            doc.add_heading(line_stripped[2:], level=1)
        elif line_stripped.startswith("- ") or line_stripped.startswith("* "):
            p = doc.add_paragraph(line_stripped[2:], style="List Bullet")
        elif line_stripped.startswith("```"):
            continue  # skip code fences
        elif line_stripped:
            doc.add_paragraph(line_stripped)
        else:
            doc.add_paragraph("")  # blank line

    # 输出到 BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{request.title or 'evogen-export'}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
